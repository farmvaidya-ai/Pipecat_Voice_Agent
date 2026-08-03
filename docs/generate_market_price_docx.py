"""One-off script: renders docs/market_price_implementation.md as a .docx.

Run: python docs/generate_market_price_docx.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn

OUT_PATH = Path(__file__).parent / "market_price_implementation.docx"

CODE_FONT = "Consolas"
BODY_FONT = "Calibri"
ACCENT = RGBColor(0x1F, 0x5C, 0x3D)  # dark green, agri-themed
DARK_TEXT = RGBColor(0x22, 0x22, 0x22)


def set_code_block_shading(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "F2F2F2",
    })
    pPr.append(shd)


def add_code_block(doc, text):
    for line in text.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        set_code_block_shading(p)
        run = p.add_run(line if line.strip() else " ")
        run.font.name = CODE_FONT
        run.font.size = Pt(9.5)
        run.font.color.rgb = DARK_TEXT


def add_body(doc, text, bold_bits=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(10.5)
    return p


def add_bullet(doc, label, rest):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    if label:
        r1 = p.add_run(label)
        r1.bold = True
        r1.font.name = CODE_FONT
        r1.font.size = Pt(10)
        r1.font.color.rgb = ACCENT
    r2 = p.add_run(rest)
    r2.font.name = BODY_FONT
    r2.font.size = Pt(10.5)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT
        run.font.name = BODY_FONT
    return h


def build():
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK_TEXT

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_heading("Market Price (Mandi) Lookup — Implementation", level=0)
    for run in title.runs:
        run.font.color.rgb = ACCENT
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sr = subtitle.add_run(
        "How Farm Vaidya answers “what's the price of X in Y” calls: where the "
        "data comes from, how it's stored, how a live call resolves it, and how it "
        "reaches the LLM as a tool call."
    )
    sr.italic = True
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Overview
    add_heading(doc, "Overview", 1)
    add_body(
        doc,
        "One pipeline scrapes Agmarknet (India's government mandi price portal) for "
        "every commodity it tracks across every state/UT, caches the results, and "
        "serves them to callers through a single LLM tool, get_price. There is no "
        "separate “arrival quantity” system — price and arrival data are scraped "
        "and stored together, one row per (state, district, commodity).",
    )
    add_code_block(
        doc,
        "agmarknet.gov.in  ->  agmarknet_scraper.py  ->  last_known_prices.json\n"
        "                                                      |\n"
        "                                          price_lookup.py (fuzzy match,\n"
        "                                          live fallback, LLM tool)\n"
        "                                                      |\n"
        "                                              Bot.py registers get_price\n"
        "                                                      |\n"
        "                                                caller hears an answer",
    )

    # 1. Scraping
    add_heading(doc, "1. Scraping — bot_processors/agmarknet_scraper.py", 1)
    add_body(
        doc,
        "Drives Agmarknet's “Daily Price and Arrival Report” with Playwright "
        "(headless Chromium).",
    )
    add_bullet(
        doc, "scrape_all(resume=False) — ",
        "the one entry point that produces the live cache. For every commodity in "
        "agmarknet_commodities.json (604 raw names, ~585 unique after normalization), "
        "it queries all states/UTs at once in “Both” mode (price + arrival columns "
        "in a single query) rather than one query per state. resume=True skips "
        "commodities that already have today's date in the cache, for restarting a "
        "partial run.",
    )
    add_bullet(
        doc, "scrape_single(state, group, commodity, crop_keyword) — ",
        "one commodity, one state. Used only for the live-fallback path (below), not "
        "the bulk run.",
    )
    add_bullet(
        doc, "_row_to_result(row) — ",
        "converts one raw scraped row (Commodity, Market, District, State/UT, Price "
        "Date, Min/Max/Modal Price, Price Unit, Arrival Quantity, Arrival Unit) into "
        "the cache's normalized schema: prices converted to per-kg (Price Unit is "
        "quintal/kg), arrival_qty/arrival_unit only included when the site actually "
        "reported one (NR or blank is dropped). Rows with a price unit other than "
        "quintal/kg (e.g. Rs./Unit, Rs./Bundle) are skipped — there's no reliable "
        "per-kg conversion for those.",
    )
    add_bullet(
        doc, "Captcha — ",
        "Agmarknet started requiring one on every “Go” submission (2026-07-31). "
        "Solved via OCR (_solve_captcha, _get_ocr_reader — a lazy singleton, the "
        "model load is the expensive part) rather than an external captcha-solving "
        "service, since the image is a clean 6-character alphanumeric render.",
    )
    add_bullet(
        doc, "harvest_commodity_reference() / harvest_state_reference() — ",
        "one-off scrapes of the form's own dropdown options, persisted to "
        "agmarknet_commodities.json / agmarknet_states.json. These are the reference "
        "lists price_lookup.py fuzzy-matches a caller's spoken commodity/state "
        "against — not hand-maintained.",
    )
    add_bullet(
        doc, "save_scraped(scraped) — ",
        "merges (not replaces) into last_known_prices.json: reads the existing file, "
        "dict.update()s in the new keys, writes the whole file back. A key this run "
        "didn't touch (e.g. a crop/state combo that came back empty today) is left "
        "as-is.",
    )
    add_body(
        doc,
        "CLI: python -m bot_processors.agmarknet_scraper (add --resume, "
        "--harvest-commodities, or --harvest-states).",
    )

    # 2. Storage
    add_heading(doc, "2. Storage — last_known_prices.json / bot_processors/price_shared.py", 1)
    add_body(
        doc,
        "A flat JSON object, one entry per “state||district||crop_keyword” key "
        "(_LAST_KNOWN_KEY_SEP = \"||\"). crop_keyword_for() normalizes a commodity name "
        "the same way on both the write side (scraper) and read side (lookup), so a "
        "live call's resolved commodity always matches what the scraper wrote:",
    )
    add_code_block(
        doc,
        "def normalize_commodity_name(name): return name.split(\"(\")[0].strip().lower()\n"
        "def crop_keyword_for(name): return normalize_commodity_name(name)"
        ".replace(\" \", \"_\").replace(\"/\", \"_\")",
    )
    add_body(doc, "Each entry:")
    add_code_block(
        doc,
        '"Karnataka||Haveri||green_chilli": {\n'
        '  "commodity": "Green Chilli", "market": "Ranebennur APMC",\n'
        '  "district": "Haveri", "state": "Karnataka",\n'
        '  "arrival_date": "02-08-2026",\n'
        '  "modal_per_kg": 20.0, "min_per_kg": 10.0, "max_per_kg": 25.0,\n'
        '  "stale": false, "arrival_qty": "2.00", "arrival_unit": "Metric Tonnes"\n'
        '}',
    )
    add_body(
        doc,
        "stale is not trusted from the stored file — it's recomputed on every read "
        "(see below), so it always reflects whether arrival_date is actually today.",
    )
    add_body(
        doc,
        "Not an append-only history: it's a rolling “latest known per key” cache. "
        "Writes fully overwrite the file each time (not atomic — no temp-file/rename "
        "swap), so a crash mid-write can in principle corrupt it, though this hasn't "
        "been observed in practice.",
    )

    # 3. Lookup
    add_heading(doc, "3. Lookup — bot_processors/price_lookup.py", 1)
    add_body(doc, "This is what the LLM actually calls.")
    add_bullet(
        doc, "_resolve_commodity(term) / _resolve_state(term) — ",
        "fuzzy-match (difflib, via _fuzzy_match) the LLM's extracted commodity/state "
        "against agmarknet_commodities.json / agmarknet_states.json. These reference "
        "files are reloaded automatically if their mtime changes "
        "(_reload_commodity_reference_if_changed, _reload_state_reference_if_changed), "
        "so a --harvest-* re-run is picked up without restarting the bot.",
    )
    add_bullet(
        doc, "fetch_price(crop_keyword, (state, district)) — ",
        "the exact-key cache read. Recomputes stale = arrival_date != today on every "
        "call.",
    )
    add_bullet(
        doc, "fetch_price_for_district(...) — ",
        "the real entry point. District names are never hand-maintained (Agmarknet's "
        "form only takes a State; “All Districts” returns every district's rows in "
        "one query), so this fuzzy-matches the caller's named district against "
        "whatever districts are already cached for that (state, crop). Only on a "
        "genuine miss does it pay for _live_fetch_commodity() -> scrape_single() — a "
        "real-time one-commodity/one-state scrape — then retries the district match "
        "against the fresh result. _inflight_live_fetches dedupes concurrent callers "
        "asking about the same uncached (state, crop) so they share one browser run "
        "instead of each launching their own.",
    )
    add_bullet(
        doc, "_merge_into_last_known(scraped) — ",
        "folds a live-fallback result into both the in-memory cache (so the current "
        "call can use it immediately) and disk (via save_scraped, so every other "
        "district in that state benefits going forward too).",
    )
    add_bullet(
        doc, "make_get_price(serializer) — ",
        "builds the actual get_price tool function. Docstring is the tool "
        "description the LLM sees; it explicitly tells the model: only call once "
        "commodity + state + district are all known, and infer the state yourself "
        "from a well-known city/district rather than asking. Logs outcome via "
        "call_db.log_tool_call.",
    )

    # 4. Bot integration
    add_heading(doc, "4. Bot integration — Bot.py", 1)
    add_code_block(
        doc,
        "get_weather = make_get_weather(serializer)\n"
        "get_price = make_get_price(serializer)\n"
        "context = LLMContext(messages, tools=[get_weather, get_price])",
    )
    add_body(
        doc,
        "get_price/get_weather are real LLM function-calling tools registered on the "
        "context — not a pipeline FrameProcessor pre-injecting data. The model "
        "decides when to call them based on their docstrings and the conversation so "
        "far.",
    )

    # 5. RAG gating
    add_heading(doc, "5. RAG gating — bot_processors/intent_router.py + bot_processors/rag.py", 1)
    add_body(
        doc,
        "classify_intent(text) is a cheap, synchronous, purely keyword-based check "
        "(Telugu/Hindi/Tamil/Kannada/English keyword sets for “weather” and "
        "“price”) run on every interim STT fragment and the final turn text. It "
        "exists only to decide whether rag.py's RAGInjector should skip its "
        "knowledge-base search for a turn — it is explicitly not what decides whether "
        "get_price/get_weather get called; that's the LLM's own judgment.",
    )
    add_body(
        doc,
        "When RAG finds nothing (_RAG_EMPTY_MARKER), a system message is injected "
        "telling the model to stay in character rather than answer as a generic AI "
        "— otherwise identity-style questions (“who made you?”) sometimes broke "
        "character with no injected context to anchor on.",
    )

    bug = doc.add_paragraph()
    bug.paragraph_format.space_before = Pt(6)
    bug.paragraph_format.space_after = Pt(6)
    set_code_block_shading(bug)
    br = bug.add_run("Bug fixed 2026-08-03")
    br.bold = True
    br.font.name = BODY_FONT
    br.font.size = Pt(10.5)
    br.font.color.rgb = RGBColor(0xB0, 0x30, 0x20)
    br2 = bug.add_run(
        " (call_919390427476_244b1ab7.log): a caller asked for green chilli price in "
        "Haveri using casual phrasing with no ധర/రేటు keyword (“గ్రీన్ చిల్లీ "
        "ఎంత చెప్తావా?”). classify_intent returned “other”, RAG ran and found "
        "nothing relevant, and injected the empty-KB marker right as the caller "
        "finished giving commodity+state+district. The model never called get_price "
        "at all for that turn — confirmed via the log, it was invoked exactly once "
        "in the whole call, for tomato — and instead hallucinated “data not "
        "available,” even though last_known_prices.json already had the exact entry. "
        "Two fixes: intent_router now recognizes this colloquial "
        "“ఎంత చెప్తావా/చెప్పు/ఉంది” pattern as price intent, and the "
        "empty-KB marker now explicitly says an empty KB result is not a reason to "
        "skip get_price/get_weather when the model already has enough to call them."
    )
    br2.font.name = BODY_FONT
    br2.font.size = Pt(10.5)
    br2.font.color.rgb = DARK_TEXT

    # 6. Scheduled refresh
    add_heading(doc, "6. Scheduled refresh — bot_processors/scraper_daemon.py", 1)
    add_body(
        doc,
        "A plain long-running loop (replaces a Windows Task Scheduler entry that "
        "kept getting auto-disabled, likely flagged by antivirus heuristics for "
        "launching a headless browser). Every 8 hours (3x/day): "
        "save_scraped(scrape_all()) then re-exports the .xlsx. Run with "
        "python -m bot_processors.scraper_daemon; stop by closing its console or "
        "taskkill /PID <pid> /F.",
    )

    # 7. Export
    add_heading(doc, "7. Human-readable export — bot_processors/export_prices_xlsx.py", 1)
    add_body(
        doc,
        "export_to_xlsx(output_path, arrival_date=None) dumps last_known_prices.json "
        "to an .xlsx (columns: state, district, commodity, market, arrival_date, "
        "modal_per_kg, min_per_kg, max_per_kg, stale, arrival_qty, arrival_unit) for "
        "manual review/filtering in Excel. Purely a read-side convenience — doesn't "
        "affect what the bot itself reads/writes. --date=DD-MM-YYYY filters to rows "
        "still carrying that arrival_date.",
    )

    # Current state
    add_heading(doc, "Current data state (as of 2026-08-03)", 1)
    add_body(
        doc,
        "last_known_prices.json holds 3,525 entries, all dated 01-08-2026 or "
        "02-08-2026 — deliberately rebuilt from an earlier nationwide scrape run so "
        "the bot answers with Saturday/Sunday data rather than today's, with no "
        "today-dated entries mixed in.",
    )

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    build()
